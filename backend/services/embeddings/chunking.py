"""文本切块与父子文档结构构建。

切块策略：
- 项目代码文件：按 2000 字符切、重叠 200，因为 Agent 需要整段代码上下文；
- 知识库文档：按 1400 字符切、重叠 150（约 500~800 token），便于精确定位答案；
- 父子检索：每个子块携带父文本（整篇文档，超长时截断），命中子块后可用
  父文本替换子块喂给 LLM，避免“命中片段但上下文不足”。
"""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path

from backend.services.embeddings.store import ChunkRecord

LOGGER = logging.getLogger(__name__)

PROJECT_CHUNK_CHARS = 2000
PROJECT_CHUNK_OVERLAP = 200
KNOWLEDGE_CHUNK_CHARS = 1400
KNOWLEDGE_CHUNK_OVERLAP = 150
MAX_PARENT_TEXT_CHARS = 20_000


def split_text(text: str, *, max_chars: int, overlap: int) -> list[str]:
    """按字符数切块，支持相邻块重叠；空白文本返回空列表。"""

    normalized = (text or "").strip()
    if not normalized:
        return []
    max_chars = max(1, min(max_chars, 20_000))
    overlap = max(0, min(overlap, max_chars // 2))
    if len(normalized) <= max_chars:
        return [normalized]

    pieces: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + max_chars, len(normalized))
        piece = normalized[start:end].strip()
        if piece:
            pieces.append(piece)
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return pieces


def _chunk_id(scope: str, source_path: str, index: int, position: str = "") -> str:
    """生成确定性的块 ID，便于增量重建时按来源整体替换。"""

    digest = hashlib.sha256(f"{scope}\n{source_path}\n{index}\n{position}".encode()).hexdigest()[
        :24
    ]
    return f"chunk_{digest}"


def build_chunks(
    *,
    scope: str,
    source_type: str,
    source_path: str,
    text: str,
    model: str = "",
    max_chars: int = KNOWLEDGE_CHUNK_CHARS,
    overlap: int = KNOWLEDGE_CHUNK_OVERLAP,
    position: str = "",
    index_offset: int = 0,
) -> list[ChunkRecord]:
    """把一段文本切成带父子结构的向量块记录。

    ``parent_text`` 默认保存整篇文本（截断到上限），供父子检索使用。
    """

    pieces = split_text(text, max_chars=max_chars, overlap=overlap)
    parent_text = text.strip()[:MAX_PARENT_TEXT_CHARS]
    return [
        ChunkRecord(
            chunk_id=_chunk_id(scope, source_path, index_offset + index, position),
            chunk_index=index_offset + index,
            chunk_text=piece,
            embedding=[],
            model=model,
            parent_id=source_path,
            parent_text=parent_text,
            position=position,
        )
        for index, piece in enumerate(pieces)
    ]


def extract_pdf_pages(path: Path) -> list[str]:
    """按页抽取 PDF 文本，返回每页一个字符串；失败时返回空列表。"""

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 - 单文档解析失败不应中断整个知识库
        LOGGER.warning("解析 PDF 失败：%s（%s）", path, exc)
        return []


def extract_document_text(path: Path) -> str:
    """按扩展名解析知识库文档为纯文本（支持 md/txt/pdf/docx）。"""

    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".txt", ".markdown"}:
            return path.read_text("utf-8", errors="replace")
        if suffix == ".pdf":
            return _extract_pdf_text(path)
        if suffix == ".docx":
            return _extract_docx_text(path)
    except OSError as exc:
        LOGGER.warning("读取知识库文档失败：%s（%s）", path, exc)
    return ""


def _extract_pdf_text(path: Path) -> str:
    """使用 pypdf 抽取 PDF 每页文本；失败时返回空字符串。"""

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted.strip():
                pages.append(extracted.strip())
        return "\n\n".join(pages)
    except Exception as exc:  # noqa: BLE001 - 单文档解析失败不应中断整个知识库
        LOGGER.warning("解析 PDF 失败：%s（%s）", path, exc)
        return ""


def _extract_docx_text(path: Path) -> str:
    """使用 python-docx 抽取段落与表格文本。"""

    try:
        import docx

        document = docx.Document(str(path))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001 - 单文档解析失败不应中断整个知识库
        LOGGER.warning("解析 DOCX 失败：%s（%s）", path, exc)
        return ""
