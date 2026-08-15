"""向量检索基础设施单元测试：切块、存储、客户端与降级。"""

from __future__ import annotations

import json
from pathlib import Path

import httpx
import pytest

from backend.core.config import get_settings
from backend.services.embeddings.chunking import build_chunks, split_text
from backend.services.embeddings.jina_client import JinaClient, JinaError
from backend.services.embeddings.retrieval import (
    _rrf_merge,
    hybrid_search_project,
    search_knowledge,
)
from backend.services.embeddings.store import (
    ChunkRecord,
    get_usage_totals,
    record_usage,
    search_vectors,
    upsert_chunks,
)
from backend.services.workspace.database import initialize_database


def _isolated_db(monkeypatch, tmp_path: Path) -> None:
    """把数据目录指向临时目录并清空 Settings 缓存。"""

    monkeypatch.setenv("AGENT_DATA_DIR", str(tmp_path / "data"))
    get_settings.cache_clear()


def test_split_text_respects_max_chars_and_overlap() -> None:
    """切块应在超长文本上生成多块并保留重叠。"""

    text = "".join(f"块{i:02d}" for i in range(30))
    pieces = split_text(text, max_chars=24, overlap=4)
    assert len(pieces) > 1
    assert all(0 < len(piece) <= 24 for piece in pieces)
    # 相邻块之间存在重叠（中文按字符计）。
    assert pieces[0][-2:] in pieces[1]


def test_build_chunks_keeps_parent_text() -> None:
    """父子检索结构：子块携带父文本与来源路径。"""

    chunks = build_chunks(
        scope="knowledge",
        source_type="doc",
        source_path="docs/manual.md",
        text="第一章 简介\n" + "内容" * 500,
        max_chars=80,
        overlap=10,
    )
    assert len(chunks) > 1
    assert chunks[0].parent_id == "docs/manual.md"
    assert "第一章 简介" in chunks[0].parent_text
    assert all(chunk.chunk_text for chunk in chunks)


def test_rrf_merge_prefers_common_items() -> None:
    """RRF 融合应让同时出现在两个列表里的键获得更高分数。"""

    merged = _rrf_merge([["a", "b", "c"], ["b", "c", "d"]])
    assert merged["b"] > merged["a"]
    assert merged["b"] > merged["d"]


@pytest.mark.asyncio
async def test_upsert_and_search_vectors(monkeypatch, tmp_path: Path) -> None:
    """向量写入后可被余弦相似度召回，且来源替换语义正确。"""

    _isolated_db(monkeypatch, tmp_path)
    await initialize_database()
    chunks = [
        ChunkRecord(
            chunk_id=f"c{index}",
            chunk_index=index,
            chunk_text=f"文本 {index}",
            embedding=vector,
            position=f"第{index + 1}页" if index == 0 else "",
        )
        for index, vector in enumerate(([1, 0, 0], [0, 1, 0], [0, 0, 1]))
    ]
    await upsert_chunks(
        scope="knowledge",
        source_type="doc",
        source_path="manual.md",
        chunks=chunks,
        model="test-model",
    )
    hits = await search_vectors(scope="knowledge", query_vector=[0.9, 0.1, 0.0], limit=2)
    assert hits[0]["sourcePath"] == "manual.md"
    assert hits[0]["chunkIndex"] == 0
    assert hits[0]["score"] > 0.9
    assert hits[0]["position"] == "第1页"

    # 重新写入同一来源应整体替换，不残留旧块。
    await upsert_chunks(
        scope="knowledge",
        source_type="doc",
        source_path="manual.md",
        chunks=[chunks[1]],
        model="test-model",
    )
    hits_after = await search_vectors(scope="knowledge", query_vector=[1, 0, 0], limit=5)
    assert len(hits_after) == 1


@pytest.mark.asyncio
async def test_usage_totals_accumulate(monkeypatch, tmp_path: Path) -> None:
    """Token 用量统计应按模型与操作累计。"""

    _isolated_db(monkeypatch, tmp_path)
    await initialize_database()
    await record_usage(model="m1", operation="embed", prompt_tokens=10, total_tokens=12)
    await record_usage(model="m1", operation="embed", prompt_tokens=8, total_tokens=9)
    totals = await get_usage_totals()
    assert totals["totalTokens"] == 21
    assert totals["items"][0]["totalTokens"] == 21


@pytest.mark.asyncio
async def test_jina_client_embed_and_rerank_use_transport() -> None:
    """Jina 客户端应使用注入的 transport 并正确解析响应。"""

    seen_tasks: list[str] = []

    def handler(request: httpx.Request) -> httpx.Response:
        payload = json.loads(request.content)
        if request.url.path.endswith("/embeddings"):
            seen_tasks.append(str(payload.get("task") or ""))
            return httpx.Response(
                200,
                json={
                    "model": payload["model"],
                    "data": [{"index": 0, "embedding": [1.0, 2.0, 3.0]}],
                    "usage": {"prompt_tokens": 3, "total_tokens": 4},
                },
            )
        return httpx.Response(
            200,
            json={
                "model": payload["model"],
                "results": [{"index": 0, "relevance_score": 0.9, "document": "doc"}],
            },
        )

    transport = httpx.MockTransport(handler)
    client = JinaClient(
        "test-key",
        embedding_model="embed-test",
        rerank_model="rerank-test",
        transport=transport,
        timeout_seconds=10,
    )
    # 默认 task 应为索引侧 retrieval.passage，查询侧应显式传 retrieval.query。
    vectors, usage = await client.embed_texts(["hello"])
    assert vectors == [[1.0, 2.0, 3.0]]
    assert usage.total_tokens == 4
    query_vectors, _query_usage = await client.embed_texts(["query"], task="retrieval.query")
    assert query_vectors == [[1.0, 2.0, 3.0]]
    assert seen_tasks == ["retrieval.passage", "retrieval.query"]

    reranked = await client.rerank("query", ["doc"], top_n=1)
    assert reranked[0]["index"] == 0
    assert reranked[0]["score"] == 0.9


def test_jina_client_requires_api_key() -> None:
    """缺少 API Key 时客户端应抛出可识别异常。"""

    get_settings.cache_clear()
    with pytest.raises(JinaError):
        JinaClient("")


@pytest.mark.asyncio
async def test_jina_client_splits_batches_and_rate_limits() -> None:
    """客户端应按 token 上限拆批，并在接近分钟限额时等待窗口滚动。"""

    request_count = 0

    def handler(request: httpx.Request) -> httpx.Response:
        nonlocal request_count
        request_count += 1
        payload = json.loads(request.content)
        items = list(payload["input"])
        return httpx.Response(
            200,
            json={
                "model": payload["model"],
                "data": [{"index": index, "embedding": [1.0]} for index in range(len(items))],
                "usage": {
                    "prompt_tokens": 10 * len(items),
                    "total_tokens": 10 * len(items),
                },
            },
        )

    transport = httpx.MockTransport(handler)
    client = JinaClient(
        "test-key",
        embedding_model="embed-test",
        transport=transport,
        timeout_seconds=10,
        max_batch=64,
        tokens_per_minute=25,
        max_call_tokens=15,
        window_seconds=0.2,
    )
    # 每个文本估算约 13 token（10 字符 //2 + 8），超过 max_call_tokens 时各自成批。
    texts = ["0123456789", "abcdefghij", "ABCDEFGHIJ"]
    vectors, usage = await client.embed_texts(texts)
    assert len(vectors) == 3
    assert request_count == 3
    assert usage.total_tokens == 30


def test_extract_docx_text_reads_paragraphs_and_tables() -> None:
    """python-docx 应能抽取段落与表格文本。"""

    import io

    import docx

    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("第一条段落")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "单元格A"
    table.rows[0].cells[1].text = "单元格B"
    document.save(buffer)

    target = Path("sample.docx")
    try:
        target.write_bytes(buffer.getvalue())
        from backend.services.embeddings.chunking import extract_document_text

        text = extract_document_text(target)
        assert "第一条段落" in text
        assert "单元格A | 单元格B" in text
    finally:
        target.unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_hybrid_search_falls_back_to_keyword_when_disabled(
    monkeypatch, tmp_path: Path
) -> None:
    """Jina 关闭或缺少密钥时，项目检索应降级为关键词并正常返回。"""

    _isolated_db(monkeypatch, tmp_path)
    monkeypatch.setenv("JINA_EMBEDDING_ENABLED", "0")
    get_settings.cache_clear()
    await initialize_database()
    results = await hybrid_search_project("project_missing", "登录超时")
    assert isinstance(results, list)


@pytest.mark.asyncio
async def test_search_knowledge_returns_empty_without_key(monkeypatch, tmp_path: Path) -> None:
    """未配置 Jina 密钥时知识库检索应返回空列表而不是报错。"""

    _isolated_db(monkeypatch, tmp_path)
    monkeypatch.delenv("JINA_API_KEY", raising=False)
    get_settings.cache_clear()
    await initialize_database()
    result = await search_knowledge("知识库问题")
    assert result.sources == []
    assert result.candidate_count == 0
